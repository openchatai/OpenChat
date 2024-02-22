# views.py
from docx import Document
from openpyxl import load_workbook
import csv
import json
import io
import os
import requests
import traceback
from pypdfium2 import PdfDocument
from uuid import uuid4
import logging.config
from typing import Optional
from django.views.decorators.csrf import csrf_exempt
from django.utils import timezone
from django.shortcuts import get_object_or_404
from django.conf import settings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders.directory import DirectoryLoader
from langchain_community.document_loaders import PyPDFium2Loader
from langchain_community.document_loaders import TextLoader
from api.utils import get_embeddings
from api.utils import init_vector_store
from api.interfaces import StoreOptions
from api.utils.make_chain import process_text_with_llm
from web.utils.delete_foler import delete_folder
from web.models.failed_jobs import FailedJob
from web.models.pdf_data_sources import PdfDataSource
from typing import Optional, Dict, Any, List

logging.config.dictConfig(settings.LOGGING)
logger = logging.getLogger(__name__)


@csrf_exempt
def pdf_handler(
    shared_folder: str,
    namespace: str,
    metadata: Dict[str, Any],
    delete_folder_flag: Optional[bool] = False,
    ocr_pdf_file: Optional[bool] = False,
    text_data: Optional[str] = None,
):
    """
    This function handles PDF files and other types of files in a shared folder. It processes the text data if provided directly,
    otherwise it reads from the files in the shared folder. It processes each file based on its extension, converts .doc, .docx, .xls,
    and .xlsx files to .txt, and saves .txt, .csv, and .json files as .txt. It then converts the text data to a vector database.

    Args:
        shared_folder (str): The name of the shared folder where the files are located.
        namespace (str): The namespace for the vector database.
        delete_folder_flag (bool): A flag indicating whether to delete the folder after processing the files.
        ocr_pdf_file (bool): A flag indicating whether to send or not PDF to OCR API services.
        text_data (Optional[str], optional): The text data to be processed. If this is provided, the function will not read from
        the files. Defaults to None.

    Raises:
        Exception: If an error occurs during the processing of the files or the conversion of the text data to a vector database.
    """

    # If text data is provided directly, process it without reading from files
    if text_data:
        process_text_data(text_data, namespace)
        logging.debug(
            "Debug: text_data is provided directly, process it without reading from files"
        )
        return

    # Convert delete_folder_flag and ocr_pdf_file to boolean (send 0 - FALSE or 1 - TRUE)
    delete_folder_flag = (
        bool(delete_folder_flag) if delete_folder_flag is not None else False
    )
    ocr_pdf_file = bool(ocr_pdf_file) if ocr_pdf_file is not None else False
    logging.debug(
        f"Debug: delete_folder_flag: {delete_folder_flag}, ocr_pdf_file: {ocr_pdf_file}"
    )

    # Check if the shared_folder is provided, if not, return early as there are no files to process
    if not shared_folder:
        print("No shared folder provided for file processing.")
        return

    try:
        # TODO: When will be multiple external library to choose, need to change.
        directory_path = os.path.join("website_data_sources", shared_folder)
        logging.debug(f"Debug: Processing folder {directory_path}")

        if not os.path.exists(directory_path):
            print(f"Debug: Directory {directory_path} does not exist.")
            return

        logging.debug(f"Debug: Directory exists. Files: {os.listdir(directory_path)}")

        # Process each file in the directory based on its extension
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)
            if filename.endswith(".pdf"):
                if ocr_pdf_file == True:
                    process_pdf(file_path, directory_path)
                    print(f"Debug: OCR PDF file {ocr_pdf_file}")
                else:
                    process_pdf_with_pypdfium(file_path, directory_path)
                    print(f"Debug: Not need to send to OCR API {ocr_pdf_file}")
            elif filename.endswith((".txt", ".csv", ".json")):
                save_as_txt(file_path)
            elif filename.endswith((".doc", ".docx", ".xls", ".xlsx")):
                convert_to_txt(file_path)

        txt_to_vectordb(shared_folder, namespace, delete_folder_flag, metadata)

    except Exception as e:
        # pdf_data_source.ingest_status = 'failed'
        # pdf_data_source.save()
        failed_job = FailedJob(
            uuid=str(uuid4()),
            connection="default",
            queue="default",
            payload="pdf_handler",
            exception=str(e),
            failed_at=timezone.now(),
        )
        failed_job.save()
        print("Exception occurred:", e)
        traceback.print_exc()


@csrf_exempt
def process_pdf_with_pypdfium(file_path, directory_path):
    pdf_document = PdfDocument(file_path)
    text_pages_with_numbers = []

    for page_index in range(len(pdf_document)):
        page = pdf_document.get_page(page_index)
        text_page = page.get_textpage()  # get a text page handle for this page
        text = text_page.get_text_range()  # extract text from the text page
        text_pages_with_numbers.append(
            (page_index + 1, text)
        )  # Store page number and text
        text_page.close()  # close the text page handle

    # Combine texts from all pages, prepending each with its page number
    combined_text = "\n".join(
        [f"Page {num}: {text}" for num, text in text_pages_with_numbers]
    )
    txt_file_path = os.path.splitext(file_path)[0] + ".txt"
    logging.debug(
        f"Debug: Writing text with page numbers to {txt_file_path}, directory_path: {directory_path}"
    )

    with open(txt_file_path, "w") as f:
        f.write(combined_text)

    pdf_document.close()


@csrf_exempt
def process_pdf(FilePath, directory_path):
    UserName = os.environ.get("OCR_USERNAME")
    LicenseCode = os.environ.get("OCR_LICCODE")
    gettext = True
    outputformat = "txt"
    language = os.environ.get("OCR_LANGUAGE", "english")
    pagerange = "allpages"
    resturl = "http://www.ocrwebservice.com/restservices/processDocument"

    RequestUrl = f"{resturl}?pagerange={pagerange}&language={language}&outputformat={outputformat}&gettext={gettext}"
    logging.debug(f"Debug: RequestUrl: {RequestUrl}")

    try:
        with open(FilePath, "rb") as image_file:
            image_data = image_file.read()
    except FileNotFoundError:
        #   pdf_data_source.ingest_status = 'failed'
        #   pdf_data_source.save()
        failed_job = FailedJob(
            uuid=str(uuid4()),
            connection="default",
            queue="default",
            payload=FilePath,
            exception="File not found",
            failed_at=timezone.now(),
        )
        failed_job.save()
        print(f"File not found: {FilePath}")
        return

    try:
        r = requests.post(RequestUrl, data=image_data, auth=(UserName, LicenseCode))

        # Decode Output response
        jobj = json.loads(r.content)

        ocrError = str(jobj["ErrorMessage"])

        if ocrError != "":
            # Error occurs during recognition
            raise Exception(f"Recognition Error:  {ocrError}")

        # Extracted text from first or single page
        # print(str(jobj["OCRText"]))

        # Extracted text from first or single page
        ocrText = str(jobj["OCRText"])

        # Extract the filename without the extension
        file_path = os.path.splitext(os.path.basename(FilePath))[0]

        # Create a new TXT file with the same name in the same directory
        txt_file_path = os.path.join(directory_path, file_path + ".txt")
        if os.environ.get("OCR_LLM", "0") == "1":
            # Write the OCR text into an in-memory text stream
            txt_file = io.StringIO(ocrText)
            mode = "assistant"
            initial_prompt = (
                f"Objective: To enhance official documents written. "
                f"\nInput Data: The text of a document which may contain grammatical errors, typos, formatting issues, and stylistic inconsistencies from OCR result. "
                f"\nFunctional Requirements: Detection and Correction of Grammatical and Typographical Errors: Identify and correct spelling and punctuation errors. Check grammatical agreements within sentences."
                f"\nStandardization of Style: Adjust the text to ensure coherence and stylistic uniformity in accordance with official writing standards."
                f"\nClarification of Text Structure: Restructure sentences to improve clarity and readability, without altering the original meaning. Keep and answer the detected language from the document."
                f"\nDocument Formatting: Implement a formatting system that adjusts the alignment of text, lists, and other structural elements for a professional presentation."
                f"\nOutput Data: This is the corrected and enhanced document. Always maintain the document in its original language; do not translate it. Respond only in the language detected from the document. Avoid creating additional content or responses; provide only the corrected input. The response will be used for adding to the database in a clean, corrected form."
                f"\nThe text: {{text}}. "
            )

            logging.debug(f"Debug: initial_prompt: {initial_prompt}")

            # Call LLM and write the result into a new text file
            process_text_with_llm(txt_file, mode, initial_prompt)
            final_text = txt_file.getvalue()
            with open(txt_file_path, "w") as f:
                f.write(final_text)
        else:
            # Write the OCR text into the new TXT file
            with open(txt_file_path, "w") as txt_file:
                txt_file.write(ocrText)

    except Exception as e:
        if str(e) == "Recognition Error: Maximum page limit exceeded":
            print("The document exceeds the maximum page limit for the OCR service.")
        else:
            print(f"Exception occurred: {e}")
        # pdf_data_source.ingest_status = 'failed'
        # pdf_data_source.save()
        failed_job = FailedJob(
            uuid=str(uuid4()),
            connection="default",
            queue="default",
            payload=FilePath,
            exception=str(e),
            failed_at=timezone.now(),
        )
        failed_job.save()
        traceback.print_exc()


@csrf_exempt
def txt_to_vectordb(
    shared_folder: str,
    namespace: str,
    delete_folder_flag: bool,
    metadata: Optional[Dict[str, Any]] = None,
):
    try:
        pdf_data_source = PdfDataSource.objects.get(folder_name=shared_folder)
        directory_path = os.path.join("website_data_sources", shared_folder)

        if metadata is None:
            metadata = {}

        # TODO: When will be multiple external library to choose, need to change.
        if os.environ.get("PDF_LIBRARY") == "external":
            directory_loader = DirectoryLoader(
                directory_path,
                glob="**/*.txt",
                loader_cls=TextLoader,
                use_multithreading=True,
            )
        else:
            directory_loader = DirectoryLoader(
                directory_path,
                glob="**/*.pdf",
                loader_cls=PyPDFium2Loader,
                use_multithreading=True,
            )

        raw_docs = directory_loader.load()

        for doc in raw_docs:
            doc.metadata = (
                getattr(doc, "metadata", {})
                if getattr(doc, "metadata", {}) is not None
                else {}
            )

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000, chunk_overlap=200, length_function=len
        )

        docs = text_splitter.split_documents(raw_docs)
        logging.debug("external files docs -->", docs)

        embeddings = get_embeddings()

        logging.debug(
            f"Initializing vector store for namespace: {namespace} with {len(docs)} documents."
        )
        init_vector_store(
            docs,
            embeddings,
            StoreOptions(namespace=namespace),
            metadata={
                "bot_id": str(pdf_data_source.chatbot.id),
                "last_update": pdf_data_source.updated_at.strftime("%Y-%m-%d %H:%M:%S"),
                "type": "document",
                "doc_type": (
                    pdf_data_source.files_info[0]["original_name"].split(".")[-1]
                    if pdf_data_source.files_info
                    else "unknown"
                ),
                "page": "1",  # @TODO to extract the page number.
                "folder": pdf_data_source.folder_name,
                "original_filename": (
                    pdf_data_source.files_info[0]["original_name"]
                    if pdf_data_source.files_info
                    else "unknown"
                ),
            },
        )
        logging.debug(
            f"Vector store initialized successfully for metadata: {metadata}."
        )

        logging.debug(f"Folder need or not to delete. {delete_folder_flag}")
        # Delete folder if flag is set
        if delete_folder_flag:
            delete_folder(folder_path=directory_path)
            logging.debug(f"All is done, folder deleted {delete_folder_flag}")

    except Exception as e:
        # pdf_data_source.ingest_status = 'failed'
        # pdf_data_source.save()
        failed_job = FailedJob(
            uuid=str(uuid4()),
            connection="default",
            queue="default",
            payload="txt_to_vectordb",
            exception=str(e),
            failed_at=timezone.now(),
        )
        failed_job.save()
        print(
            f"Failed to initialize vector store for namespace: {namespace}. Exception: {e}"
        )
        traceback.print_exc()


def save_as_txt(file_path):
    txt_file_path = os.path.splitext(file_path)[0] + ".txt"
    with open(file_path, "rb") as file:
        with open(txt_file_path, "wb") as txt_file:
            txt_file.write(file.read())


def convert_to_txt(file_path):
    txt_file_path = os.path.splitext(file_path)[0] + ".txt"
    if file_path.endswith(".docx"):
        document = Document(file_path)
        with open(txt_file_path, "w", encoding="utf-8") as txt_file:
            for para in document.paragraphs:
                txt_file.write(para.text + "\n")
    elif file_path.endswith(".xls"):
        # xlrd is used for .xls files, but it is not listed in the installed packages.
        # If xlrd is installed, the following code can be used:
        raise NotImplementedError("Conversion from .xls to text not implemented yet.")
    elif file_path.endswith(".xlsx"):
        workbook = load_workbook(filename=file_path)
        with open(txt_file_path, "w", encoding="utf-8") as txt_file:
            for sheet in workbook:
                for row in sheet.iter_rows(values_only=True):
                    txt_file.write(
                        ",".join(
                            [str(cell) if cell is not None else "" for cell in row]
                        )
                        + "\n"
                    )
    elif file_path.endswith(".csv"):
        with open(file_path, "r", encoding="utf-8") as csv_file:
            reader = csv.reader(csv_file)
            with open(txt_file_path, "w", encoding="utf-8") as txt_file:
                for row in reader:
                    txt_file.write(",".join(row) + "\n")
    elif file_path.endswith(".json"):
        with open(file_path, "r", encoding="utf-8") as json_file:
            data = json.load(json_file)
            with open(txt_file_path, "w", encoding="utf-8") as txt_file:
                json.dump(data, txt_file, ensure_ascii=False, indent=4)
    else:
        raise NotImplementedError(
            f"Conversion for {os.path.splitext(file_path)[1]} files to text not implemented yet."
        )


def process_text_data(text_data: str, namespace: str):
    """
    Processes the provided text data and ingests it into the vector database.

    Args:
        text_data (str): The text data to process.
        namespace (str): The namespace for the vector database.
    """
    logging.debug("Debug: process_text_data")
    txt_to_vectordb(text_data, namespace, False)
    return
